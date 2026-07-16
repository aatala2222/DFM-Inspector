"""
Create a Word document with roll forming tolerance information
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ROLL FORMING\nSTANDARD MANUFACTURING TOLERANCES")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Industry Standard Guidelines")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph()

# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
run.font.size = Pt(11)
run.font.italic = True

doc.add_page_break()

# Section 1: Cross-Section Tolerances
doc.add_heading('1. Cross-Section Tolerances', level=1)

doc.add_paragraph(
    'Cross-section tolerances define the allowable variation in the dimensions '
    'of the formed profile perpendicular to the length of the part.'
)

table1 = doc.add_table(rows=5, cols=2)
table1.style = 'Light Grid Accent 1'

data1 = [
    ('Dimension Type', 'Standard Tolerance'),
    ('Fractional dimensions', '±0.031" (±0.79 mm)'),
    ('Decimal dimensions', '±0.010" (±0.25 mm)'),
    ('Angular dimensions', '±1° to 2°'),
    ('Standard cross-section', '±1/64" (±0.40 mm)')
]

for i, (label, value) in enumerate(data1):
    table1.rows[i].cells[0].text = label
    table1.rows[i].cells[1].text = value
    if i == 0:
        table1.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table1.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run(
    'Under specific controlled conditions, tighter tolerances of ±0.005" (±0.13 mm) '
    'can be achieved, but this typically requires additional tooling costs, extended '
    'tryout time, and premium materials with special thickness and mechanical property controls.'
)

doc.add_page_break()

# Section 2: Length Tolerances
doc.add_heading('2. Length Tolerances', level=1)

doc.add_paragraph(
    'Length tolerances vary based on both the material thickness and the overall '
    'length of the formed part.'
)

doc.add_heading('For Thicker Parts (≥0.026" thickness)', level=2)

table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Light Grid Accent 1'

data2 = [
    ('Part Length', 'Tolerance'),
    ('Up to 36" long', '±0.015" (±0.38 mm)'),
    ('36" to 96" long', '±0.030" (±0.76 mm)'),
    ('96" to 144" long', '±0.060" (±1.52 mm)')
]

for i, (label, value) in enumerate(data2):
    table2.rows[i].cells[0].text = label
    table2.rows[i].cells[1].text = value
    if i == 0:
        table2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table2.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

doc.add_heading('For Thinner Parts (0.015" to 0.025" thickness)', level=2)

table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Light Grid Accent 1'

data3 = [
    ('Part Length', 'Tolerance'),
    ('Up to 36" long', '±0.020" (±0.51 mm)'),
    ('36" to 96" long', '±0.047" (±1.19 mm)'),
    ('96" to 144" long', '±0.093" (±2.36 mm)')
]

for i, (label, value) in enumerate(data3):
    table3.rows[i].cells[0].text = label
    table3.rows[i].cells[1].text = value
    if i == 0:
        table3.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table3.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# Section 3: Straightness Tolerances
doc.add_heading('3. Straightness Tolerances', level=1)

doc.add_paragraph(
    'Straightness tolerances account for deviations in the formed part along its length. '
    'These include camber (vertical deviation), sweep (horizontal deviation), and twist.'
)

table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Light Grid Accent 1'

data4 = [
    ('Type', 'Description', 'Standard Tolerance'),
    ('Camber/Bow', 'Vertical plane deviation', '0.015" max per foot of length'),
    ('Sweep', 'Horizontal plane deviation', '1/8" to 1/4" in 10 feet'),
    ('Twist', 'Rotational deviation', '0.5° max per foot (5° to 15° in 10 feet)')
]

for i, (type_val, desc, tol) in enumerate(data4):
    table4.rows[i].cells[0].text = type_val
    table4.rows[i].cells[1].text = desc
    table4.rows[i].cells[2].text = tol
    if i == 0:
        for j in range(3):
            table4.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# Section 4: Factors Affecting Tolerances
doc.add_heading('4. Key Factors Affecting Tolerances', level=1)

doc.add_paragraph(
    'The achievable tolerances in roll forming are influenced by multiple factors '
    'that must be considered during design and production:'
)

factors = [
    ('Material Properties', 'Yield strength, elastic modulus, and hardness variations affect springback and dimensional stability'),
    ('Material Thickness', 'Thickness variations within a coil or between coils directly impact cross-section dimensions'),
    ('Strip Width Consistency', 'Width variations affect leg lengths and overall profile dimensions'),
    ('Deformation Amount', 'Greater deformation requires more forming stages and can reduce achievable tolerance'),
    ('Tool Contact', 'Complete tool contact with all surfaces enables tighter tolerance control'),
    ('Springback', 'Elastic recovery after forming causes dimensional changes that must be compensated'),
    ('End Flare', 'Residual stresses cause greater distortion at part ends and cut edges'),
    ('Process Variables', 'Roll pressure, lubricant, temperature changes affect consistency')
]

for title, description in factors:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_page_break()

# Section 5: Design Recommendations
doc.add_heading('5. Design Recommendations for Optimal Tolerances', level=1)

doc.add_paragraph(
    'Following these design guidelines will help maintain standard tolerances and '
    'minimize manufacturing costs:'
)

recommendations = [
    ('Maximize Bend Radius', 
     'Use the maximum bend radius permissible for your application. An inside bend radius '
     'less than the material thickness will reduce roll life, increase power requirements, '
     'and make tight tolerances more difficult to achieve.'),
    
    ('Design for Symmetry', 
     'Create symmetrical parts whenever possible to eliminate twist in the finished part. '
     'Asymmetrical designs create unbalanced forces that lead to distortion.'),
    
    ('Position Features Carefully', 
     'Keep holes, slots, and notches away from bend lines to prevent distortion. Features '
     'placed directly on or too close to bends will be deformed during the forming process.'),
    
    ('Specify Realistic Tolerances', 
     'Only specify tolerances that are tighter than standard when absolutely necessary. '
     'Unnecessarily tight tolerances greatly increase both tooling costs and per-part costs.'),
    
    ('Account for Multiple Gauges', 
     'If the same part will be produced in multiple material thicknesses, understand that '
     'using one set of rolls for all gauges will result in dimensional variations between '
     'thicknesses.'),
    
    ('Consider Material Variations', 
     'Recognize that strip width, thickness, and mechanical property variations will affect '
     'the final part dimensions. Design with appropriate tolerance bands to accommodate this.'),
    
    ('Plan for Springback Compensation', 
     'Work with your roll forming supplier to over-form the part to compensate for springback. '
     'The amount of compensation depends on material properties and bend angles.'),
    
    ('Address End Flare', 
     'Understand that end flare cannot be completely eliminated. If critical dimensions exist '
     'at part ends, discuss secondary operations or design modifications with your supplier.')
]

for i, (title, description) in enumerate(recommendations, 1):
    doc.add_heading(f'{i}. {title}', level=2)
    doc.add_paragraph(description)

doc.add_page_break()

# Section 6: Springback and End Flare
doc.add_heading('6. Springback and End Flare Considerations', level=1)

doc.add_heading('Springback', level=2)
doc.add_paragraph(
    'Springback is the elastic distortion that occurs after a part is removed from forming '
    'pressure. The amount of springback varies with material properties such as yield point '
    'and elastic modulus. Roll forming strains are more complex than simple bending, creating '
    'residual stresses throughout the part.'
)

p = doc.add_paragraph()
p.add_run('Compensation Strategy: ').bold = True
p.add_run(
    'Tool designers typically over-form the part to compensate for expected springback. '
    'The amount of over-forming is determined through experience and testing with specific '
    'material grades and thicknesses.'
)

doc.add_paragraph()

doc.add_heading('End Flare', level=2)
doc.add_paragraph(
    'End flare is a distortion that occurs at the ends of roll formed sections or at any '
    'point where the cross-section is cut (such as pierces or notches). Residual stresses '
    'make themselves particularly apparent by causing greater distortion at part ends than '
    'along the length.'
)

p = doc.add_paragraph()
p.add_run('Mitigation Options: ').bold = True
p.add_run(
    'End flare can be minimized through careful roll design procedures, but it cannot be '
    'completely eliminated except by subjecting the metal to stretch forming or stress '
    'relieving anneal operations.'
)

doc.add_page_break()

# Section 7: When Tighter Tolerances Are Needed
doc.add_heading('7. Achieving Tighter Tolerances', level=1)

doc.add_paragraph(
    'When applications require tolerances tighter than standard, several approaches can be used:'
)

approaches = [
    'Premium Materials: Use materials with tighter thickness and mechanical property controls',
    'Additional Tooling: Design specialized tooling with more forming stations and tighter clearances',
    'Extended Tryout: Allow additional time for tool adjustment and process optimization',
    'Secondary Operations: Add post-forming operations such as sizing, straightening, or machining',
    'In-Process Monitoring: Implement measurement systems to detect and correct variations',
    'Environmental Controls: Maintain consistent temperature and lubrication conditions'
]

for approach in approaches:
    doc.add_paragraph(approach, style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Cost Impact: ').bold = True
run = p.add_run(
    'Tighter tolerances typically increase costs by 20-50% or more, depending on the '
    'specific requirements. Always discuss tolerance requirements early in the design '
    'process with your roll forming supplier.'
)
run.font.color.rgb = RGBColor(204, 0, 0)

doc.add_page_break()

# Section 8: Summary Table
doc.add_heading('8. Quick Reference Summary', level=1)

summary_table = doc.add_table(rows=11, cols=2)
summary_table.style = 'Medium Grid 1 Accent 1'

summary_data = [
    ('Tolerance Category', 'Standard Value'),
    ('Cross-Section (Fractional)', '±0.031"'),
    ('Cross-Section (Decimal)', '±0.010"'),
    ('Cross-Section (Standard)', '±1/64"'),
    ('Angular', '±1° to 2°'),
    ('Length (up to 36", thick)', '±0.015"'),
    ('Length (36-96", thick)', '±0.030"'),
    ('Length (96-144", thick)', '±0.060"'),
    ('Camber/Bow', '0.015" per foot'),
    ('Twist', '0.5° per foot'),
    ('Tighter Tolerances', '±0.005" (special conditions)')
]

for i, (category, value) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = category
    summary_table.rows[i].cells[1].text = value
    if i == 0:
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        summary_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# Important Note
note_p = doc.add_paragraph()
note_p.add_run('Important: ').bold = True
note_p.add_run(
    'These are general industry guidelines. Actual achievable tolerances depend on specific '
    'part geometry, material properties, and manufacturing conditions. Always discuss your '
    'tolerance requirements with your roll forming supplier during the design phase.'
)

doc.add_page_break()

# Footer section
doc.add_paragraph('_' * 80)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Roll Forming Tolerance Reference Document')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

sources = doc.add_paragraph()
sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sources.add_run(
    'Content compiled from industry standards and manufacturer guidelines\n'
    'Sources: Samson Roll Form, Formtek Group, and industry best practices'
)
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

date_footer = doc.add_paragraph()
date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_footer.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

# Save document
filename = f'Roll_Forming_Tolerances_{datetime.now().strftime("%Y-%m-%d")}.docx'
doc.save(filename)

print(f"✓ Document created successfully: {filename}")
