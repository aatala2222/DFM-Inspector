"""
Word Document Report Generator for DFM Analysis
Generates professional Word documents with analysis results and visual annotations
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, List
import os

try:
    from .visual_annotator import VisualAnnotator
    VISUAL_ANNOTATIONS_AVAILABLE = True
except ImportError:
    VISUAL_ANNOTATIONS_AVAILABLE = False
    print("Warning: Visual annotations not available. Install matplotlib to enable.")

try:
    from .cad_visualizer import CADVisualizer
    CAD_VISUALIZATION_AVAILABLE = True
except ImportError:
    CAD_VISUALIZATION_AVAILABLE = False
    print("Warning: 3D CAD visualization not available.")


class WordReportGenerator:
    """Generate Word document reports for DFM analysis"""
    
    def __init__(self, step_file_path: str = None, parser=None):
        """
        Initialize Word report generator
        
        Args:
            step_file_path: Path to STEP file (optional if parser provided)
            parser: Pre-loaded parser instance with accurate geometry
        """
        self.doc = Document()
        self._setup_styles()
        self.annotator = VisualAnnotator() if VISUAL_ANNOTATIONS_AVAILABLE else None
        self.cad_visualizer = None
        self.parser = parser
        
        if CAD_VISUALIZATION_AVAILABLE:
            try:
                # Prefer using parser if provided (more accurate)
                if parser:
                    print("✓ Creating CADVisualizer with enhanced parser")
                    self.cad_visualizer = CADVisualizer(step_file_path=step_file_path, parser=parser)
                elif step_file_path:
                    print("✓ Creating CADVisualizer with STEP file path")
                    self.cad_visualizer = CADVisualizer(step_file_path=step_file_path)
            except Exception as e:
                print(f"Could not load 3D visualization: {e}")
        
        self.generated_images = []
    
    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Title style
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading styles already exist, just customize them
        heading1 = styles['Heading 1']
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        heading2 = styles['Heading 2']
        heading2.font.color.rgb = RGBColor(0, 102, 204)
    
    def generate_report(self, analysis_results: Dict, filename: str = None) -> str:
        """
        Generate a complete DFM analysis report with visual annotations
        
        Args:
            analysis_results: Dictionary containing analysis results
            filename: Output filename (optional)
            
        Returns:
            Path to generated Word document
        """
        # Add title page
        self._add_title_page(analysis_results)
        
        # Add executive summary with visual chart
        self._add_executive_summary(analysis_results)
        
        # Add geometry analysis
        self._add_geometry_section(analysis_results)
        
        # Add visual annotations section (NEW)
        if self.annotator:
            self._add_visual_annotations_section(analysis_results)
        
        # Add rule-by-rule analysis
        self._add_rules_section(analysis_results)
        
        # Add cost optimization opportunities
        self._add_cost_optimization_section(analysis_results)
        
        # Add recommendations
        self._add_recommendations_section(analysis_results)
        
        # Save document
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            process = analysis_results.get('process', 'DFM').replace(' ', '_')
            filename = f"DFM_Report_{process}_{timestamp}.docx"
        
        self.doc.save(filename)
        
        # Cleanup temporary images
        if self.annotator:
            self.annotator.cleanup()
        
        return filename
    
    def _add_title_page(self, results: Dict):
        """Add title page with header information"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DFM ANALYSIS REPORT")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{results.get('process', 'Manufacturing Process')} Analysis")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)
        
        self.doc.add_paragraph()  # Spacing
        
        # Report information table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Report details
        details = [
            ('Date:', datetime.now().strftime("%B %d, %Y")),
            ('Process:', results.get('process', 'N/A')),
            ('Material:', results.get('material', 'N/A')),
            ('Score:', f"{results.get('score', 0)}/100"),
            ('Status:', self._get_status_text(results.get('score', 0)))
        ]
        
        for i, (label, value) in enumerate(details):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[i].cells[1].text = value
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self, results: Dict):
        """Add executive summary section"""
        self.doc.add_heading('Executive Summary', level=1)
        
        # Overall assessment
        score = results.get('score', 0)
        assessment = self._get_status_text(score)
        
        p = self.doc.add_paragraph()
        p.add_run('Overall Assessment: ').bold = True
        run = p.add_run(assessment)
        run.font.color.rgb = self._get_status_color(score)
        run.font.bold = True
        
        # Score
        p = self.doc.add_paragraph()
        p.add_run('Manufacturability Score: ').bold = True
        p.add_run(f"{score}/100")
        
        # Summary statistics
        self.doc.add_paragraph()
        stats_table = self.doc.add_table(rows=4, cols=2)
        stats_table.style = 'Light List Accent 1'
        
        stats = [
            ('Critical Issues', results.get('issues', 0)),
            ('Warnings', results.get('warnings', 0)),
            ('Passed Checks', results.get('passed', 0)),
            ('Optimization Opportunities', results.get('suggestions', 0))
        ]
        
        for i, (label, value) in enumerate(stats):
            stats_table.rows[i].cells[0].text = label
            stats_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            stats_table.rows[i].cells[1].text = str(value)
        
        # Summary text
        if 'summary' in results:
            self.doc.add_paragraph()
            summary_text = results['summary'].replace('**', '').replace('•', '  •')
            self.doc.add_paragraph(summary_text)
        
        self.doc.add_page_break()
    
    def _add_visual_annotations_section(self, results: Dict):
        """Add visual annotations showing problem areas with 3D renderings"""
        print("\n" + "="*70)
        print("VISUAL ANNOTATIONS SECTION")
        print("="*70)
        
        self.doc.add_heading('Visual Analysis - Problem Areas', level=1)
        
        use_3d = self.cad_visualizer is not None
        print(f"3D Visualizer available: {use_3d}")
        
        # Check if this is an enhanced DFM workflow result with violations
        has_violations = 'violations' in results and 'visualization_data' in results
        
        if has_violations:
            print("✓ Enhanced DFM workflow detected - using violation-based visualization")
            self._add_enhanced_violations_visualization(results)
            return
        
        # Original visualization logic for legacy results
        if use_3d:
            self.doc.add_paragraph(
                'The following 3D renderings show the actual CAD geometry with problematic '
                'features highlighted in red. These images are generated directly from your '
                'STEP file to show exactly where design changes are needed.'
            )
        else:
            self.doc.add_paragraph(
                'The following images highlight specific features that failed DFM rules. '
                'Red indicates failures, orange indicates warnings, and green indicates passed checks.'
            )
        
        # Generate summary chart
        if self.annotator:
            summary_img = self.annotator.create_summary_chart(results)
            if summary_img and os.path.exists(summary_img):
                self.doc.add_paragraph()
                self.doc.add_picture(summary_img, width=Inches(6))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Get geometry and holes data
        geometry = results.get('geometry', {})  # Raw geometry data with dict dimensions
        geometry_info = results.get('geometry_info', {})  # Formatted strings for display
        all_rules = results.get('all_rules', [])
        
        # Try to extract holes from multiple sources
        holes = []
        
        # First, try to get holes from parser's analysis dict
        if 'parser' in results:
            parser = results['parser']
            print(f"Parser found in results: {type(parser)}")
            if hasattr(parser, 'analysis') and isinstance(parser.analysis, dict):
                holes = parser.analysis.get('holes', [])
                print(f"✓ Found {len(holes)} holes from parser.analysis")
            elif hasattr(parser, 'holes'):
                holes = parser.holes
                print(f"✓ Found {len(holes)} holes from parser.holes")
            else:
                print(f"⚠ Parser has no 'analysis' dict or 'holes' attribute")
                print(f"Parser attributes: {dir(parser)}")
        else:
            print("⚠ No parser in results")
        
        # Fallback: try to get from results directly
        if not holes and 'holes' in results:
            holes = results.get('holes', [])
            print(f"✓ Found {len(holes)} holes from results")
        
        print(f"Total holes available for visualization: {len(holes)}")
        print(f"All rules count: {len(all_rules)}")
        print(f"Geometry info: {geometry}")
        print("="*70 + "\n")
        
        # If we have a 3D visualizer and failed rules, generate at least one overview rendering
        failed_rules = [r for r in all_rules if r.get('status') in ['FAIL', 'WARNING']]
        images_added = 0
        
        if use_3d and failed_rules:
            print(f"\n→ Generating overview 3D rendering for {len(failed_rules)} failed rules")
            try:
                # Use CADVisualizer's professional rendering
                import matplotlib
                matplotlib.use('Agg')
                from matplotlib import pyplot as plt
                from mpl_toolkits.mplot3d.art3d import Poly3DCollection
                import tempfile
                
                fig = plt.figure(figsize=(14, 11), facecolor='white')
                ax = fig.add_subplot(111, projection='3d', facecolor='white')
                
                # Render the part with CAD-quality styling
                if self.cad_visualizer.vertices is not None and self.cad_visualizer.faces is not None:
                    poly = Poly3DCollection(
                        self.cad_visualizer.vertices[self.cad_visualizer.faces],
                        alpha=0.95,
                        facecolors='#B8C5D6',  # Light blue-gray CAD color
                        edgecolors='#8A9AAA',  # Subtle gray edge lines to show features
                        linewidths=0.15,  # Very thin lines
                        antialiased=True
                    )
                    ax.add_collection3d(poly)
                    
                    # Set axis properties with CAD styling
                    self.cad_visualizer._set_axis_properties_cad_style(ax, (30, 45))
                    
                    # Add title
                    title = f'DFM Analysis - {len(failed_rules)} Issue(s) Detected\n'
                    title += f'Critical: {len([r for r in failed_rules if r.get("status") == "FAIL"])}, '
                    title += f'Warnings: {len([r for r in failed_rules if r.get("status") == "WARNING"])}'
                    ax.set_title(title, fontsize=15, weight='bold', pad=25, color='#2C3E50')
                    
                    # Save with high quality
                    img_path = os.path.join(tempfile.gettempdir(), 'dfm_overview_3d.png')
                    plt.tight_layout()
                    plt.savefig(img_path, dpi=200, bbox_inches='tight', facecolor='white',
                               edgecolor='none', pad_inches=0.2)
                    plt.close()
                    
                    if os.path.exists(img_path):
                        print(f"  ✓ Overview 3D rendering successful: {img_path}")
                        self.doc.add_paragraph()
                        self.doc.add_picture(img_path, width=Inches(6.5))
                        last_paragraph = self.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        caption = self.doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run(
                            f'Figure: 3D CAD Model - Part Geometry\n'
                            f'Analyzed for {len(failed_rules)} DFM issue(s)'
                        )
                        run.font.size = Pt(10)
                        run.font.italic = True
                        
                        images_added = 1
                        
                    # NOW ADD ANNOTATED VERSION WITH PROBLEM AREAS HIGHLIGHTED
                    print(f"\n→ Generating annotated rendering with problem areas highlighted")
                    self._add_annotated_rendering(failed_rules, geometry)
                    images_added += 1
                    
                else:
                    print(f"  ✗ No mesh data available for rendering")
            except Exception as e:
                print(f"  ✗ Overview rendering failed: {e}")
                import traceback
                traceback.print_exc()
        
        # Generate annotations for failed rules
        for rule in all_rules:
            if rule.get('status') in ['FAIL', 'WARNING']:
                rule_name = rule.get('name', '')
                print(f"\nProcessing rule: {rule_name} (status: {rule.get('status')})")
                
                # Handle hole-related rules with 3D visualization
                if 'Hole' in rule_name and holes:
                    print(f"  → Hole-related rule with {len(holes)} holes available")
                    self.doc.add_paragraph()
                    self.doc.add_heading(f'Issue: {rule_name}', level=2)
                    
                    # Determine which holes failed
                    failed_holes = self._identify_failed_holes(rule, holes, geometry)
                    print(f"  → Identified {len(failed_holes)} failed holes")
                    
                    if failed_holes:
                        img_path = None
                        
                        # Try 3D rendering first
                        if use_3d:
                            print(f"  → Attempting 3D rendering...")
                            try:
                                img_path = self.cad_visualizer.render_with_highlighted_holes(
                                    holes, failed_holes, rule_name
                                )
                                if img_path and os.path.exists(img_path):
                                    print(f"  ✓ 3D rendering successful: {img_path}")
                                    self.doc.add_picture(img_path, width=Inches(6))
                                    last_paragraph = self.doc.paragraphs[-1]
                                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Add caption
                                    caption = self.doc.add_paragraph()
                                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = caption.add_run(
                                        f'Figure: 3D View - {rule_name}\n'
                                        f'{len(failed_holes)} failed hole(s) highlighted in red on actual part geometry'
                                    )
                                    run.font.size = Pt(10)
                                    run.font.italic = True
                                    
                                    images_added += 1
                                    
                                    # Add multiple views
                                    multiview_img = self.cad_visualizer.render_multiple_views(
                                        holes, failed_holes, rule_name
                                    )
                                    if multiview_img and os.path.exists(multiview_img):
                                        print(f"  ✓ Multi-view rendering successful: {multiview_img}")
                                        self.doc.add_paragraph()
                                        self.doc.add_picture(multiview_img, width=Inches(6.5))
                                        last_paragraph = self.doc.paragraphs[-1]
                                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        
                                        caption = self.doc.add_paragraph()
                                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        run = caption.add_run(
                                            f'Figure: Multiple Views - {rule_name}'
                                        )
                                        run.font.size = Pt(10)
                                        run.font.italic = True
                                        images_added += 1
                                else:
                                    print(f"  ✗ 3D rendering failed - no image generated")
                            except Exception as e:
                                print(f"  ✗ 3D rendering failed: {e}")
                                import traceback
                                traceback.print_exc()
                                use_3d = False  # Fall back to 2D
                        
                        # Fallback to 2D schematic if 3D failed or not available
                        if not use_3d and self.annotator:
                            print(f"  → Falling back to 2D schematic")
                            img_path = self.annotator.create_hole_annotation(
                                geometry, holes, failed_holes, rule_name
                            )
    
    def _add_enhanced_violations_visualization(self, results: Dict):
        """
        Add visualization section for enhanced DFM workflow results
        
        Uses violation data with 3D coordinates to render color-coded highlights
        """
        print("Adding enhanced violations visualization...")
        
        self.doc.add_paragraph(
            'The following 3D renderings show DFM violations highlighted directly on your part geometry. '
            'Red indicates critical issues, orange shows warnings, and yellow marks suggestions. '
            'Each violation is shown at its exact 3D location with measured values.'
        )
        
        violations_data = results.get('violations', {})
        viz_data = results.get('visualization_data', {})
        
        total_violations = violations_data.get('total_violations', 0)
        by_severity = violations_data.get('by_severity', {})
        
        # Add summary
        summary = self.doc.add_paragraph()
        summary.add_run(f"Total Violations Found: {total_violations}\n").bold = True
        summary.add_run(f"  • Critical: {by_severity.get('critical', 0)}\n")
        summary.add_run(f"  • Warnings: {by_severity.get('warning', 0)}\n")
        summary.add_run(f"  • Suggestions: {by_severity.get('suggestion', 0)}\n")
        
        if total_violations == 0:
            self.doc.add_paragraph("✓ No DFM violations detected. Design is ready for manufacturing!")
            return
        
        # Render all violations on one view
        if self.cad_visualizer and viz_data:
            try:
                # Convert visualization data to list format
                all_violations = []
                for severity in ['critical', 'warning', 'suggestion']:
                    for v in viz_data.get(severity, []):
                        all_violations.append({
                            'location': v['location'],
                            'severity': severity,
                            'feature_type': v['feature_type'],
                            'measured_value': v['measured_value'],
                            'required_value': v['required_value'],
                            'message': v['message']
                        })
                
                if all_violations:
                    print(f"Rendering {len(all_violations)} violations...")
                    
                    # Single view with all violations
                    img_path = self.cad_visualizer.render_with_violations(
                        all_violations,
                        rule_name="All DFM Violations",
                        view_angle=(30, 45)
                    )
                    
                    if img_path and os.path.exists(img_path):
                        print(f"✓ Violations rendering successful: {img_path}")
                        self.doc.add_paragraph()
                        self.doc.add_picture(img_path, width=Inches(6.5))
                        last_paragraph = self.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        caption = self.doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run(
                            f'Figure: DFM Violations Highlighted on Part Geometry\n'
                            f'Red = Critical, Orange = Warning, Yellow = Suggestion'
                        )
                        run.font.size = Pt(10)
                        run.font.italic = True
                    
                    # Multi-view rendering
                    multiview_img = self.cad_visualizer.render_violations_multiview(
                        all_violations,
                        rule_name="DFM Violations"
                    )
                    
                    if multiview_img and os.path.exists(multiview_img):
                        print(f"✓ Multi-view rendering successful: {multiview_img}")
                        self.doc.add_paragraph()
                        self.doc.add_picture(multiview_img, width=Inches(6.5))
                        last_paragraph = self.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        caption = self.doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run('Figure: Multiple Views - All Violations')
                        run.font.size = Pt(10)
                        run.font.italic = True
                
            except Exception as e:
                print(f"Error rendering violations: {e}")
                import traceback
                traceback.print_exc()
        
        # Add detailed violation list
        self.doc.add_page_break()
        self.doc.add_heading('Detailed Violation List', level=2)
        
        for violation in violations_data.get('violations', []):
            rule_name = violation.get('rule_name', 'Unknown Rule')
            severity = violation.get('severity', 'warning')
            message = violation.get('message', '')
            recommendation = violation.get('recommendation', '')
            
            # Add violation entry
            p = self.doc.add_paragraph()
            
            # Severity indicator
            severity_symbol = {'critical': '🔴', 'warning': '🟠', 'suggestion': '🟡'}.get(severity, '⚪')
            p.add_run(f"{severity_symbol} {rule_name}\n").bold = True
            p.add_run(f"Severity: {severity.upper()}\n")
            p.add_run(f"Issue: {message}\n")
            p.add_run(f"Recommendation: {recommendation}\n")
            
            # Add location info
            location = violation.get('location')
            if location:
                if isinstance(location, (list, tuple)):
                    loc_str = f"({location[0]:.2f}, {location[1]:.2f}, {location[2]:.2f})"
                else:
                    loc_str = f"({location.get('x', 0):.2f}, {location.get('y', 0):.2f}, {location.get('z', 0):.2f})"
                p.add_run(f"Location: {loc_str} mm\n")
            
            p.add_run(f"Measured: {violation.get('measured_value', 0):.2f}mm | ")
            p.add_run(f"Required: {violation.get('required_value', 0):.2f}mm\n")
        
        print(f"✓ Added {total_violations} violations to report")
    
    def _add_annotated_rendering(self, failed_rules: List[Dict], geometry: Dict):
        """
        Add annotated 3D rendering with problem areas highlighted in red
        
        Args:
            failed_rules: List of failed DFM rules
            geometry: Geometry information including dimensions
        """
        try:
            import matplotlib
            matplotlib.use('Agg')
            from matplotlib import pyplot as plt
            from mpl_toolkits.mplot3d.art3d import Poly3DCollection
            import numpy as np
            import tempfile
            
            fig = plt.figure(figsize=(14, 11), facecolor='white')
            ax = fig.add_subplot(111, projection='3d', facecolor='white')
            
            # Render the part
            if self.cad_visualizer.vertices is not None and self.cad_visualizer.faces is not None:
                poly = Poly3DCollection(
                    self.cad_visualizer.vertices[self.cad_visualizer.faces],
                    alpha=0.7,  # More transparent to see annotations
                    facecolors='#B8C5D6',
                    edgecolors='#8A9AAA',
                    linewidths=0.15,
                    antialiased=True
                )
                ax.add_collection3d(poly)
                
                # Get part bounds for annotation placement
                vertices = self.cad_visualizer.vertices
                bounds = [vertices.min(axis=0), vertices.max(axis=0)]
                dims = geometry.get('dimensions', {})
                
                # Annotate each failed rule
                for i, rule in enumerate(failed_rules):
                    rule_name = rule.get('name', 'Unknown')
                    status = rule.get('status', 'WARNING')
                    measured = rule.get('measured_value', '')
                    
                    # Determine color based on severity
                    color = '#FF3333' if status == 'FAIL' else '#FFA500'
                    
                    # Place annotation at strategic locations based on rule type
                    if 'Corner' in rule_name or 'Radius' in rule_name:
                        # Highlight corners - place markers at the corners of the part
                        corners = [
                            bounds[0],  # Min corner
                            bounds[1],  # Max corner
                            [bounds[0][0], bounds[1][1], bounds[0][2]],
                            [bounds[1][0], bounds[0][1], bounds[1][2]],
                        ]
                        
                        for corner in corners[:2]:  # Show 2 corners
                            # Draw sphere at corner
                            u = np.linspace(0, 2 * np.pi, 20)
                            v = np.linspace(0, np.pi, 20)
                            radius = (bounds[1][0] - bounds[0][0]) * 0.05  # 5% of part size
                            x_sphere = corner[0] + radius * np.outer(np.cos(u), np.sin(v))
                            y_sphere = corner[1] + radius * np.outer(np.sin(u), np.sin(v))
                            z_sphere = corner[2] + radius * np.outer(np.ones(np.size(u)), np.cos(v))
                            
                            ax.plot_surface(x_sphere, y_sphere, z_sphere,
                                          color=color, alpha=0.9, antialiased=True)
                            
                            # Add arrow
                            arrow_len = (bounds[1][2] - bounds[0][2]) * 0.15
                            ax.quiver(corner[0], corner[1], corner[2] + arrow_len,
                                    0, 0, -arrow_len * 0.6,
                                    color=color, arrow_length_ratio=0.3, linewidth=2.5)
                    
                    elif 'Wall' in rule_name or 'Thickness' in rule_name:
                        # Highlight thin wall areas - place markers on sides
                        center = (bounds[0] + bounds[1]) / 2
                        
                        # Draw sphere at center
                        u = np.linspace(0, 2 * np.pi, 20)
                        v = np.linspace(0, np.pi, 20)
                        radius = (bounds[1][0] - bounds[0][0]) * 0.06
                        x_sphere = center[0] + radius * np.outer(np.cos(u), np.sin(v))
                        y_sphere = center[1] + radius * np.outer(np.sin(u), np.sin(v))
                        z_sphere = center[2] + radius * np.outer(np.ones(np.size(u)), np.cos(v))
                        
                        ax.plot_surface(x_sphere, y_sphere, z_sphere,
                                      color=color, alpha=0.9, antialiased=True)
                        
                        # Add arrow
                        arrow_len = (bounds[1][2] - bounds[0][2]) * 0.2
                        ax.quiver(center[0], center[1], center[2] + arrow_len,
                                0, 0, -arrow_len * 0.6,
                                color=color, arrow_length_ratio=0.3, linewidth=2.5)
                    
                    # Add text label - positioned further away and in black
                    label_offset = (bounds[1][2] - bounds[0][2]) * 0.3 * (i + 1)  # Further away
                    label_pos = bounds[1] + np.array([0, 0, label_offset])
                    ax.text(label_pos[0], label_pos[1], label_pos[2],
                           f'{rule_name}\n{measured}',
                           color='black',  # Black text for readability
                           fontsize=11, weight='bold',
                           ha='center', va='bottom',
                           bbox=dict(boxstyle='round,pad=0.6',
                                   facecolor='white', edgecolor='black', linewidth=1.5))
                
                # Set axis properties
                self.cad_visualizer._set_axis_properties_cad_style(ax, (30, 45))
                
                # Add title
                title = f'Problem Areas Highlighted\n'
                title += f'Red = Critical Issues, Orange = Warnings'
                ax.set_title(title, fontsize=15, weight='bold', pad=25, color='#2C3E50')
                
                # Save
                img_path = os.path.join(tempfile.gettempdir(), 'dfm_annotated_3d.png')
                plt.tight_layout()
                plt.savefig(img_path, dpi=200, bbox_inches='tight', facecolor='white',
                           edgecolor='none', pad_inches=0.2)
                plt.close()
                
                if os.path.exists(img_path):
                    print(f"  ✓ Annotated rendering successful: {img_path}")
                    self.doc.add_paragraph()
                    self.doc.add_picture(img_path, width=Inches(6.5))
                    last_paragraph = self.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    caption = self.doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(
                        f'Figure: Problem Areas Highlighted on Part\n'
                        f'Red/Orange markers show locations of DFM issues'
                    )
                    run.font.size = Pt(10)
                    run.font.italic = True
                    
        except Exception as e:
            print(f"  ✗ Annotated rendering failed: {e}")
            import traceback
            traceback.print_exc()

        if images_added == 0:
            self.doc.add_paragraph(
                'No visual annotations generated. This may be because:\n'
                '• All rules passed (no failures to visualize)\n'
                '• Geometry data not available for visualization\n'
                '• Rule types don\'t support visual annotation yet'
            )
        else:
            # Add note about 3D vs 2D
            note = self.doc.add_paragraph()
            if use_3d:
                run = note.add_run(
                    f'✓ Generated {images_added} 3D rendering(s) from actual STEP file geometry. '
                    'Red highlights show exact locations of non-compliant features on your part.'
                )
            else:
                run = note.add_run(
                    f'Generated {images_added} schematic diagram(s). '
                    'For 3D renderings of actual part geometry, ensure trimesh is installed: pip install trimesh'
                )
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        
        self.doc.add_page_break()
    
    def _identify_failed_holes(self, rule: Dict, holes: List[Dict], geometry: Dict) -> List[Dict]:
        """Identify which specific holes failed a rule"""
        failed_holes = []
        rule_name = rule.get('name', '')
        
        if 'Diameter' in rule_name:
            # Check hole diameter rule
            standard = rule.get('standard', '')
            import re
            match = re.search(r'(\d+\.?\d*)\s*T', standard)
            if match:
                multiplier = float(match.group(1))
                thickness = geometry.get('material_thickness', 1.0)
                if thickness == 0:
                    thickness = float(geometry.get('min_thickness', '1').replace('mm', '').strip())
                
                min_diameter = multiplier * thickness
                
                for hole in holes:
                    diameter = hole.get('diameter', 0)
                    if diameter > 0 and diameter < min_diameter:
                        failed_holes.append(hole)
        
        elif 'Edge' in rule_name and 'Distance' in rule_name:
            # Check hole to edge distance
            standard = rule.get('standard', '')
            import re
            match = re.search(r'(\d+\.?\d*)\s*T', standard)
            if match:
                multiplier = float(match.group(1))
                thickness = geometry.get('material_thickness', 1.0)
                if thickness == 0:
                    thickness = float(geometry.get('min_thickness', '1').replace('mm', '').strip())
                
                min_distance = multiplier * thickness
                
                dims = geometry.get('dimensions', {})
                if isinstance(dims, str):
                    # Parse "X x Y x Z mm" format
                    parts = dims.split('x')
                    width = float(parts[0].strip()) if len(parts) > 0 else 100
                    height = float(parts[1].strip()) if len(parts) > 1 else 100
                else:
                    width = dims.get('x', 100)
                    height = dims.get('y', 100)
                
                for hole in holes:
                    center = hole.get('center', {})
                    x = center.get('x', 0)
                    y = center.get('y', 0)
                    diameter = hole.get('diameter', 0)
                    radius = diameter / 2
                    
                    # Calculate distances to edges
                    dist_left = x - radius
                    dist_right = width - (x + radius)
                    dist_bottom = y - radius
                    dist_top = height - (y + radius)
                    
                    min_edge_dist = min(dist_left, dist_right, dist_bottom, dist_top)
                    
                    if min_edge_dist < min_distance:
                        failed_holes.append(hole)
        
        elif 'Spacing' in rule_name:
            # Check hole spacing
            standard = rule.get('standard', '')
            import re
            match = re.search(r'(\d+\.?\d*)\s*T', standard)
            if match:
                multiplier = float(match.group(1))
                thickness = geometry.get('material_thickness', 1.0)
                if thickness == 0:
                    thickness = float(geometry.get('min_thickness', '1').replace('mm', '').strip())
                
                min_spacing = multiplier * thickness
                
                # Check all hole pairs
                for i, hole1 in enumerate(holes):
                    for hole2 in holes[i+1:]:
                        center1 = hole1.get('center', {})
                        center2 = hole2.get('center', {})
                        
                        x1, y1 = center1.get('x', 0), center1.get('y', 0)
                        x2, y2 = center2.get('x', 0), center2.get('y', 0)
                        
                        # Calculate center-to-center distance
                        center_dist = ((x2 - x1)**2 + (y2 - y1)**2)**0.5
                        
                        # Calculate edge-to-edge distance
                        edge_dist = center_dist - (hole1.get('diameter', 0)/2 + hole2.get('diameter', 0)/2)
                        
                        if edge_dist < min_spacing:
                            if hole1 not in failed_holes:
                                failed_holes.append(hole1)
                            if hole2 not in failed_holes:
                                failed_holes.append(hole2)
        
        return failed_holes
    
    def _add_geometry_section(self, results: Dict):
        """Add geometry analysis section"""
        self.doc.add_heading('Geometry Analysis', level=1)
        
        geometry_info = results.get('geometry_info', {})
        
        if geometry_info:
            table = self.doc.add_table(rows=len(geometry_info), cols=2)
            table.style = 'Light Grid Accent 1'
            
            for i, (key, value) in enumerate(geometry_info.items()):
                # Format key (convert snake_case to Title Case)
                formatted_key = key.replace('_', ' ').title() + ':'
                table.rows[i].cells[0].text = formatted_key
                table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                table.rows[i].cells[1].text = str(value)
        else:
            self.doc.add_paragraph('Geometry information not available.')
        
        self.doc.add_page_break()
    
    def _add_rules_section(self, results: Dict):
        """Add detailed rule-by-rule analysis"""
        self.doc.add_heading('Detailed Rule-by-Rule Analysis', level=1)
        
        all_rules = results.get('all_rules', [])
        
        if not all_rules:
            self.doc.add_paragraph('No detailed rules available for this analysis.')
            return
        
        for i, rule in enumerate(all_rules, 1):
            # Rule heading
            self.doc.add_heading(f"{i}. {rule.get('name', 'Rule')}", level=2)
            
            # Status badge
            status = rule.get('status', 'INFO')
            p = self.doc.add_paragraph()
            p.add_run('Status: ').bold = True
            status_run = p.add_run(status)
            status_run.font.bold = True
            status_run.font.color.rgb = self._get_rule_status_color(status)
            
            # Standard
            p = self.doc.add_paragraph()
            p.add_run('Standard: ').bold = True
            p.add_run(rule.get('standard', 'N/A'))
            
            # Measured value
            p = self.doc.add_paragraph()
            p.add_run('Measured Value: ').bold = True
            p.add_run(rule.get('measured_value', 'N/A'))
            
            # Evaluation
            self.doc.add_paragraph()
            self.doc.add_paragraph('Evaluation:', style='Heading 3')
            self.doc.add_paragraph(rule.get('evaluation', 'N/A'))
            
            # Recommendation
            self.doc.add_paragraph('Recommendation:', style='Heading 3')
            self.doc.add_paragraph(rule.get('recommendation', 'N/A'))
            
            # Rationale
            self.doc.add_paragraph('Rationale:', style='Heading 3')
            self.doc.add_paragraph(rule.get('rationale', 'N/A'))
            
            # Cost impact
            p = self.doc.add_paragraph()
            p.add_run('Cost Impact: ').bold = True
            p.add_run(rule.get('cost_impact', 'N/A'))
            
            # Add separator
            if i < len(all_rules):
                self.doc.add_paragraph('_' * 80)
        
        self.doc.add_page_break()
    
    def _add_cost_optimization_section(self, results: Dict):
        """Add cost optimization opportunities"""
        self.doc.add_heading('Cost Optimization Opportunities', level=1)
        
        details = results.get('details', {})
        cost_savings = details.get('cost_savings', [])
        
        if not cost_savings:
            self.doc.add_paragraph('No cost optimization opportunities identified.')
            return
        
        for i, opportunity in enumerate(cost_savings, 1):
            self.doc.add_heading(f"Opportunity {i}", level=2)
            
            # Opportunity description
            p = self.doc.add_paragraph()
            p.add_run('Opportunity: ').bold = True
            p.add_run(opportunity.get('opportunity', 'N/A'))
            
            # Savings
            p = self.doc.add_paragraph()
            p.add_run('Potential Savings: ').bold = True
            savings_run = p.add_run(opportunity.get('savings', 'N/A'))
            savings_run.font.color.rgb = RGBColor(0, 128, 0)
            savings_run.font.bold = True
            
            # Difficulty
            p = self.doc.add_paragraph()
            p.add_run('Implementation Difficulty: ').bold = True
            p.add_run(opportunity.get('difficulty', 'N/A'))
            
            # Rationale
            if 'rationale' in opportunity:
                self.doc.add_paragraph()
                self.doc.add_paragraph(opportunity['rationale'])
        
        self.doc.add_page_break()
    
    def _add_recommendations_section(self, results: Dict):
        """Add final recommendations"""
        self.doc.add_heading('Recommendations', level=1)
        
        details = results.get('details', {})
        
        # Critical issues
        critical_issues = details.get('critical_issues', [])
        if critical_issues:
            self.doc.add_heading('Critical Issues (Must Fix)', level=2)
            for issue in critical_issues:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{issue.get('category', 'Issue')}: ").bold = True
                p.add_run(issue.get('message', 'N/A'))
                
                # Recommendation
                rec_p = self.doc.add_paragraph(style='List Bullet 2')
                rec_p.add_run('→ Recommendation: ').italic = True
                rec_p.add_run(issue.get('recommendation', 'N/A'))
        
        # Warnings
        warnings = details.get('warnings', [])
        if warnings:
            self.doc.add_paragraph()
            self.doc.add_heading('Warnings (Should Fix)', level=2)
            for warning in warnings:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{warning.get('category', 'Warning')}: ").bold = True
                p.add_run(warning.get('message', 'N/A'))
                
                # Recommendation
                rec_p = self.doc.add_paragraph(style='List Bullet 2')
                rec_p.add_run('→ Recommendation: ').italic = True
                rec_p.add_run(warning.get('recommendation', 'N/A'))
        
        # Final recommendation
        self.doc.add_paragraph()
        self.doc.add_heading('Final Recommendation', level=2)
        
        score = results.get('score', 0)
        if score >= 90:
            recommendation = "Design is excellent and ready for manufacturing. Proceed with confidence."
        elif score >= 75:
            recommendation = "Design is good with minor improvements recommended. Address warnings to optimize cost."
        elif score >= 60:
            recommendation = "Design is acceptable but has several areas for improvement. Review all warnings before production."
        else:
            recommendation = "Design needs significant revision. Address all critical issues before proceeding to manufacturing."
        
        self.doc.add_paragraph(recommendation)
        
        # Add footer with source
        self.doc.add_paragraph()
        self.doc.add_paragraph('_' * 80)
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('Generated by DFM Inspector')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _get_status_text(self, score: float) -> str:
        """Get status text based on score"""
        if score >= 90:
            return "EXCELLENT"
        elif score >= 75:
            return "GOOD"
        elif score >= 60:
            return "ACCEPTABLE"
        else:
            return "NEEDS REVISION"
    
    def _get_status_color(self, score: float) -> RGBColor:
        """Get color based on score"""
        if score >= 90:
            return RGBColor(0, 128, 0)  # Green
        elif score >= 75:
            return RGBColor(0, 102, 204)  # Blue
        elif score >= 60:
            return RGBColor(255, 140, 0)  # Orange
        else:
            return RGBColor(204, 0, 0)  # Red
    
    def _get_rule_status_color(self, status: str) -> RGBColor:
        """Get color based on rule status"""
        status_colors = {
            'PASS': RGBColor(0, 128, 0),      # Green
            'FAIL': RGBColor(204, 0, 0),      # Red
            'WARNING': RGBColor(255, 140, 0), # Orange
            'INFO': RGBColor(0, 102, 204)     # Blue
        }
        return status_colors.get(status, RGBColor(0, 0, 0))


def generate_word_report(analysis_results: Dict, output_path: str = None) -> str:
    """
    Convenience function to generate a Word report
    
    Args:
        analysis_results: Dictionary containing analysis results
        output_path: Optional output path for the document
        
    Returns:
        Path to generated Word document
    """
    generator = WordReportGenerator()
    return generator.generate_report(analysis_results, output_path)
