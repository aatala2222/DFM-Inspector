"""
Generate DFM Rules Comprehensive Guide as a Word document.
Reads from DFM_RULES_COMPREHENSIVE_GUIDE.md and produces a formatted .docx file.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import date
import re

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# -- Style definitions --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 51, 102)

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(11)

def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = text.strip()
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            if i < cols:
                row.cells[i].text = text.strip()
                for p in row.cells[i].paragraphs:
                    p.style = doc.styles['Normal']
                    for run in p.runs:
                        run.font.size = Pt(9)
    doc.add_paragraph()  # spacing after table
    return table


def parse_md_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (header, rows, end_idx)."""
    header_line = lines[start_idx].strip()
    header = [c.strip() for c in header_line.split('|') if c.strip()]
    # skip separator line
    data_rows = []
    idx = start_idx + 2
    while idx < len(lines) and '|' in lines[idx] and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].split('|') if c.strip()]
        data_rows.append(row)
        idx += 1
    return header, data_rows, idx


# -- Read the markdown file --
with open('DFM_RULES_COMPREHENSIVE_GUIDE.md', 'r', encoding='utf-8') as f:
    content = f.read()
lines = content.split('\n')

# -- Title Page --
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DFM Rules\nComprehensive Guide')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0, 51, 102)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('All Manufacturing Processes\nDesign for Manufacturability Reference')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Generated: {date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info2.add_run('Processes: CNC Machining | Injection Molding | Die Casting & Permanent Mold\nSheet Metal | Weldments | Integral Skin PU Foam Molding')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# -- Parse and render content --
i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip the very first title block (already on title page)
    if stripped.startswith('# DFM Rules Comprehensive Guide'):
        i += 1
        continue
    if stripped.startswith('## All Manufacturing'):
        i += 1
        continue
    if stripped.startswith('**Document:**') or stripped.startswith('**Source Spec'):
        i += 1
        continue
    if stripped.startswith('**Processes Covered:**'):
        i += 1
        continue

    # Skip table of contents section
    if stripped == '## Table of Contents':
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#'):
            i += 1
        continue

    # Horizontal rule = skip
    if stripped == '---':
        i += 1
        continue

    # End-of-document note
    if stripped.startswith('*End of DFM Rules'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.strip('*'))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        i += 1
        continue

    # Heading 1: # N. Title
    if re.match(r'^# \d+\.', stripped):
        doc.add_page_break()
        doc.add_heading(stripped.lstrip('# '), level=1)
        i += 1
        continue

    # Heading 2: ## N.N Title
    if re.match(r'^## \d+\.\d+', stripped):
        doc.add_heading(stripped.lstrip('# '), level=2)
        i += 1
        continue

    # Heading 3: ### ...
    if stripped.startswith('### '):
        doc.add_heading(stripped.lstrip('# '), level=3)
        i += 1
        continue

    # Source/process info lines
    if stripped.startswith('**Source:**') or stripped.startswith('**Process Owner:**') or stripped.startswith('**Standards:**'):
        p = doc.add_paragraph()
        # Parse bold prefix
        m = re.match(r'\*\*(.+?)\*\*\s*(.*)', stripped)
        if m:
            run = p.add_run(m.group(1) + ' ')
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(m.group(2))
            run.font.size = Pt(9)
        else:
            p.add_run(stripped).font.size = Pt(9)
        i += 1
        continue

    # Table detection
    if '|' in stripped and stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
        header, data_rows, end_idx = parse_md_table(lines, i)
        add_table(doc, header, data_rows)
        i = end_idx
        continue

    # Bullet points
    if stripped.startswith('- '):
        text = stripped[2:]
        # Handle bold prefix in bullet
        p = doc.add_paragraph(style='List Bullet')
        if '**' in text:
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            p.add_run(text)
        for run in p.runs:
            run.font.size = Pt(10)
        i += 1
        continue

    # Italic lines (notes)
    if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.strip('*'))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        i += 1
        continue

    # Bold paragraph lines
    if stripped.startswith('**') and stripped.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.strip('*'))
        run.bold = True
        i += 1
        continue

    # Regular paragraph (skip empty)
    if stripped:
        p = doc.add_paragraph()
        # Handle inline bold
        if '**' in stripped:
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            p.add_run(stripped)
        i += 1
        continue

    i += 1

# -- Save --
output_file = f'DFM_Rules_Comprehensive_Guide_{date.today().strftime("%Y-%m-%d")}.docx'
doc.save(output_file)
print(f"Word document created: {output_file}")
