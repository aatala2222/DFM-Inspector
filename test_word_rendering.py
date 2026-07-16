"""
Test Word Report 3D Rendering
Verifies that 3D renderings are included in Word reports
"""
import os
import sys
from docx import Document

print("=" * 70)
print("WORD REPORT 3D RENDERING TEST")
print("=" * 70)

# Find the most recent DFM report
import glob
reports = glob.glob("DFM_Report_*.docx")
if not reports:
    print("\n✗ No DFM reports found")
    print("  Please generate a report first:")
    print("  1. Go to http://localhost:5000")
    print("  2. Upload a STEP file")
    print("  3. Click Analyze")
    print("  4. Click Export to Word")
    sys.exit(1)

# Get the most recent report
latest_report = max(reports, key=os.path.getctime)
print(f"\n✓ Found report: {latest_report}")
print(f"  Size: {os.path.getsize(latest_report):,} bytes")
print(f"  Modified: {os.path.getctime(latest_report)}")

# Open and analyze the document
try:
    doc = Document(latest_report)
    print(f"\n✓ Opened document successfully")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")
    
    # Count images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    
    print(f"  Images: {image_count}")
    
    # Check for specific sections
    sections_found = {
        'Visual Analysis': False,
        'Figure': False,
        '3D': False
    }
    
    for para in doc.paragraphs:
        text = para.text
        if 'Visual Analysis' in text:
            sections_found['Visual Analysis'] = True
        if 'Figure' in text:
            sections_found['Figure'] = True
        if '3D' in text:
            sections_found['3D'] = True
    
    print(f"\n✓ Content analysis:")
    for section, found in sections_found.items():
        status = "✓" if found else "✗"
        print(f"  {status} {section}: {'Found' if found else 'Not found'}")
    
    # Verdict
    print(f"\n" + "=" * 70)
    if image_count > 0:
        print(f"✅ SUCCESS: Report contains {image_count} image(s)")
        print(f"   3D renderings are included in the Word document!")
    else:
        print(f"❌ ISSUE: Report contains NO images")
        print(f"   3D renderings are missing from the Word document")
        print(f"\n   Possible causes:")
        print(f"   1. CADVisualizer not initialized (check trimesh installation)")
        print(f"   2. STEP file parsing failed")
        print(f"   3. No violations detected (perfect part)")
        print(f"\n   Check server console output for errors")
    print("=" * 70)
    
except Exception as e:
    print(f"\n✗ Error analyzing document: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
